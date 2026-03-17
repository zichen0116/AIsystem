import json
import logging
import re
import tempfile
import uuid
from pathlib import Path
from typing import Annotated, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from starlette.background import BackgroundTask

from app.core.auth import CurrentUser
from app.core.database import get_db
from app.models.chat_history import ChatHistory
from app.models.lesson_plan import LessonPlan
from app.models.lesson_plan_reference import LessonPlanReference
from app.schemas.lesson_plan import (
    LessonPlanExportRequest,
    LessonPlanGenerateRequest,
    LessonPlanLatestResponse,
    LessonPlanModifyRequest,
    LessonPlanSaveRequest,
    LessonPlanSaveResponse,
    LessonPlanUploadResponse,
    LessonPlanInfo,
    MessageInfo,
    FileInfo,
    LessonPlanListResponse,
    LessonPlanListItem,
    LessonPlanMessagesResponse,
    ChatMessageInfo,
)
from app.services.lesson_plan_service import (
    LESSON_PLAN_MODIFY_PROMPT,
    LESSON_PLAN_SYSTEM_PROMPT,
    get_chat_history,
    retrieve_context,
    save_after_stream,
    stream_llm,
)

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/lesson-plan", tags=["lesson-plan"])

DbSession = Annotated[AsyncSession, Depends(get_db)]


# --------------- SSE helpers ---------------

async def _sse_generate(plan_id: int, session_id: str, user_id: int, messages: list[dict]):
    """SSE generator: emit metadata, stream LLM tokens, then persist to DB."""
    yield f"data: {json.dumps({'meta': {'lesson_plan_id': plan_id, 'session_id': session_id}}, ensure_ascii=False)}\n\n"

    full = ""
    try:
        async for chunk in stream_llm(messages):
            full += chunk
            yield f"data: {json.dumps({'content': chunk}, ensure_ascii=False)}\n\n"
    except Exception as e:
        logger.error(f"SSE stream error: {e}")
        yield f"data: {json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"

    yield "data: [DONE]\n\n"
    await save_after_stream(plan_id, session_id, user_id, full)


def _sse_response(generator):
    return StreamingResponse(
        generator,
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


# --------------- 1. Generate ---------------

@router.post("/generate")
async def generate_lesson_plan(req: LessonPlanGenerateRequest, user: CurrentUser, db: DbSession):
    """Create or reuse a LessonPlan, retrieve context, stream LLM generation."""

    # 1. Create or find LessonPlan
    lesson_plan = None
    if req.session_id:
        result = await db.execute(
            select(LessonPlan).where(LessonPlan.session_id == uuid.UUID(req.session_id), LessonPlan.user_id == user.id)
        )
        lesson_plan = result.scalar_one_or_none()

    if not lesson_plan:
        lesson_plan = LessonPlan(user_id=user.id, status="generating")
        db.add(lesson_plan)
        await db.flush()

    # 2. Write user message to ChatHistory
    db.add(ChatHistory(session_id=lesson_plan.session_id, user_id=user.id, role="user", content=req.query))

    # 3. Retrieve context (RAG + reference files)
    context = await retrieve_context(db, req.query, req.library_ids, req.file_ids, user.id)

    # 4. Read existing chat history for multi-turn
    history = await get_chat_history(db, lesson_plan.session_id, limit=10)

    await db.commit()

    # 5. Build LLM messages
    messages = [{"role": "system", "content": LESSON_PLAN_SYSTEM_PROMPT}]
    for msg in history:
        messages.append({"role": msg["role"], "content": msg["content"]})
    user_content = req.query
    if context:
        user_content = f"以下是参考资料：\n\n{context}\n\n---\n\n用户要求：{req.query}"
    # Replace last user message with enriched version (context included)
    if messages and messages[-1]["role"] == "user":
        messages[-1]["content"] = user_content
    else:
        messages.append({"role": "user", "content": user_content})

    return _sse_response(_sse_generate(lesson_plan.id, str(lesson_plan.session_id), user.id, messages))


# --------------- 2. Modify ---------------

@router.post("/modify")
async def modify_lesson_plan(req: LessonPlanModifyRequest, user: CurrentUser, db: DbSession):
    """Modify existing lesson plan: read history from DB, stream LLM."""

    # Verify ownership
    result = await db.execute(
        select(LessonPlan).where(LessonPlan.id == req.lesson_plan_id, LessonPlan.user_id == user.id)
    )
    lesson_plan = result.scalar_one_or_none()
    if not lesson_plan:
        raise HTTPException(404, "教案不存在")

    # Write user instruction to ChatHistory
    db.add(ChatHistory(session_id=lesson_plan.session_id, user_id=user.id, role="user", content=req.instruction))

    # Read recent history from DB
    history = await get_chat_history(db, lesson_plan.session_id, limit=10)

    # Retrieve optional reference context
    context = await retrieve_context(db, req.instruction, req.library_ids, req.file_ids, user.id)

    await db.commit()

    # Build LLM messages
    messages = [{"role": "system", "content": LESSON_PLAN_MODIFY_PROMPT}]
    for msg in history[:-1]:  # history minus current instruction (already in prompt)
        messages.append({"role": msg["role"], "content": msg["content"]})

    user_msg = f"当前教案内容:\n\n{req.current_content}\n\n"
    if context:
        user_msg += f"参考资料：\n\n{context}\n\n"
    user_msg += f"修改要求：{req.instruction}"
    messages.append({"role": "user", "content": user_msg})

    return _sse_response(_sse_generate(lesson_plan.id, str(lesson_plan.session_id), user.id, messages))


# --------------- 3. Upload ---------------

@router.post("/upload", response_model=LessonPlanUploadResponse)
async def upload_reference_file(
    user: CurrentUser,
    db: DbSession,
    file: UploadFile = File(...),
    lesson_plan_id: Optional[int] = Form(None),
):
    """Upload a reference file, parse content, store in LessonPlanReference."""
    from app.services.parsers.factory import ParserFactory

    # Save to disk
    upload_dir = Path("uploads/lesson_plan")
    upload_dir.mkdir(parents=True, exist_ok=True)
    dest = upload_dir / f"{uuid.uuid4()}{Path(file.filename).suffix}"
    content_bytes = await file.read()
    dest.write_bytes(content_bytes)

    # Parse
    parsed = ""
    try:
        result = await ParserFactory.parse_file(str(dest))
        if result and result.chunks:
            parsed = "\n\n".join(c.content for c in result.chunks)
    except Exception as e:
        logger.warning(f"Parse failed for {file.filename}: {e}")
        parsed = content_bytes.decode("utf-8", errors="ignore")

    # Persist
    ref = LessonPlanReference(
        user_id=user.id,
        lesson_plan_id=lesson_plan_id,
        filename=file.filename,
        file_path=str(dest),
        parsed_content=parsed,
    )
    db.add(ref)
    await db.flush()

    return LessonPlanUploadResponse(file_id=ref.id, filename=ref.filename)


# --------------- 4. Save (auto-save) ---------------

@router.patch("/{lesson_plan_id}", response_model=LessonPlanSaveResponse)
async def save_lesson_plan(lesson_plan_id: int, req: LessonPlanSaveRequest, user: CurrentUser, db: DbSession):
    """Auto-save editor content (called on blur / 30s timer)."""
    result = await db.execute(
        select(LessonPlan).where(LessonPlan.id == lesson_plan_id, LessonPlan.user_id == user.id)
    )
    plan = result.scalar_one_or_none()
    if not plan:
        raise HTTPException(404, "教案不存在")

    if req.content is not None:
        plan.content = req.content
    if req.title is not None:
        plan.title = req.title
    await db.flush()

    return LessonPlanSaveResponse(id=plan.id, updated_at=plan.updated_at.isoformat() if plan.updated_at else "")


# --------------- 5. Load latest ---------------

@router.get("/latest", response_model=LessonPlanLatestResponse)
async def get_latest_lesson_plan(user: CurrentUser, db: DbSession):
    """Load the most recent lesson plan with its chat history and reference files."""
    result = await db.execute(
        select(LessonPlan)
        .where(LessonPlan.user_id == user.id)
        .order_by(LessonPlan.updated_at.desc())
        .limit(1)
    )
    plan = result.scalar_one_or_none()

    if not plan:
        return LessonPlanLatestResponse()

    # Chat history
    hist_result = await db.execute(
        select(ChatHistory)
        .where(ChatHistory.session_id == plan.session_id)
        .order_by(ChatHistory.created_at)
    )
    msgs = [MessageInfo(role=h.role, content=h.content) for h in hist_result.scalars().all()]

    # Reference files
    files_result = await db.execute(
        select(LessonPlanReference)
        .where(LessonPlanReference.lesson_plan_id == plan.id)
    )
    files = [FileInfo(id=f.id, filename=f.filename) for f in files_result.scalars().all()]

    return LessonPlanLatestResponse(
        lesson_plan=LessonPlanInfo(
            id=plan.id,
            session_id=str(plan.session_id),
            title=plan.title,
            content=plan.content,
            status=plan.status,
        ),
        messages=msgs,
        files=files,
    )


# --------------- 6. List ---------------

@router.get("/list", response_model=LessonPlanListResponse)
async def list_lesson_plans(user: CurrentUser, db: DbSession):
    """获取当前用户的所有教案列表，按创建时间倒序"""
    stmt = select(LessonPlan).where(LessonPlan.user_id == user.id).order_by(LessonPlan.created_at.desc())
    result = await db.execute(stmt)
    plans = result.scalars().all()

    items = [
        LessonPlanListItem(
            id=p.id,
            session_id=str(p.session_id),
            title=p.title,
            status=p.status,
            created_at=p.created_at.isoformat(),
            updated_at=p.updated_at.isoformat(),
        )
        for p in plans
    ]

    return LessonPlanListResponse(items=items, total=len(items))


# --------------- 7. Detail ---------------

@router.get("/{lesson_plan_id}", response_model=LessonPlanInfo)
async def get_lesson_plan_detail(lesson_plan_id: int, user: CurrentUser, db: DbSession):
    """获取指定教案的详细信息"""
    result = await db.execute(
        select(LessonPlan).where(LessonPlan.id == lesson_plan_id, LessonPlan.user_id == user.id)
    )
    plan = result.scalar_one_or_none()
    if not plan:
        raise HTTPException(404, "教案不存在")

    return LessonPlanInfo(
        id=plan.id,
        session_id=str(plan.session_id),
        title=plan.title,
        content=plan.content,
        status=plan.status,
    )


# --------------- 9. Messages ---------------

@router.get("/{lesson_plan_id}/messages", response_model=LessonPlanMessagesResponse)
async def get_lesson_plan_messages(lesson_plan_id: int, user: CurrentUser, db: DbSession):
    """获取指定教案的对话历史"""
    # 验证教案存在且属于当前用户
    plan_result = await db.execute(
        select(LessonPlan).where(LessonPlan.id == lesson_plan_id, LessonPlan.user_id == user.id)
    )
    plan = plan_result.scalar_one_or_none()
    if not plan:
        raise HTTPException(404, "教案不存在")

    # 获取对话历史
    hist_result = await db.execute(
        select(ChatHistory)
        .where(ChatHistory.session_id == plan.session_id)
        .order_by(ChatHistory.created_at)
    )
    messages = [
        ChatMessageInfo(role=h.role, content=h.content, created_at=h.created_at.isoformat())
        for h in hist_result.scalars().all()
    ]

    return LessonPlanMessagesResponse(messages=messages)


# --------------- 10. Export DOCX ---------------

def _export_docx_with_python_docx(markdown_content: str, output_path: str) -> None:
    """Best-effort Markdown -> DOCX fallback without pandoc."""
    from docx import Document

    doc = Document()
    lines = markdown_content.splitlines()

    task_re = re.compile(r"^\s*[-*]\s+\[([ xX])\]\s+(.*)$")
    bullet_re = re.compile(r"^\s*[-*]\s+(.*)$")
    ordered_re = re.compile(r"^\s*\d+\.\s+(.*)$")

    for raw in lines:
        stripped = raw.strip()
        if not stripped:
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            continue

        task_match = task_re.match(stripped)
        if task_match:
            checked = task_match.group(1).lower() == "x"
            text = task_match.group(2).strip()
            p = doc.add_paragraph(("☑ " if checked else "☐ ") + text)
            p.style = "List Bullet"
            continue

        bullet_match = bullet_re.match(stripped)
        if bullet_match:
            p = doc.add_paragraph(bullet_match.group(1).strip())
            p.style = "List Bullet"
            continue

        ordered_match = ordered_re.match(stripped)
        if ordered_match:
            p = doc.add_paragraph(ordered_match.group(1).strip())
            p.style = "List Number"
            continue

        doc.add_paragraph(stripped)

    doc.save(output_path)


@router.post("/export/docx")
async def export_docx(req: LessonPlanExportRequest, user: CurrentUser):
    """Convert Markdown to DOCX. Prefer pandoc, fallback to python-docx."""

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = tmp.name

    try:
        last_error = None

        # Primary path: pypandoc + pandoc
        try:
            import pypandoc

            try:
                pypandoc.convert_text(req.content, "docx", format="md", outputfile=tmp_path)
                last_error = None
            except OSError:
                try:
                    pypandoc.download_pandoc()
                    pypandoc.convert_text(req.content, "docx", format="md", outputfile=tmp_path)
                    last_error = None
                except Exception as e:
                    last_error = e
            except Exception as e:
                last_error = e
        except Exception as e:
            last_error = e

        # Fallback path: python-docx (no pandoc dependency)
        if last_error is not None:
            try:
                _export_docx_with_python_docx(req.content, tmp_path)
                last_error = None
            except Exception as e:
                last_error = e

        if last_error is not None:
            raise last_error

        return FileResponse(
            tmp_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{req.title}.docx",
            background=BackgroundTask(lambda: Path(tmp_path).unlink(missing_ok=True)),
        )
    except Exception as e:
        Path(tmp_path).unlink(missing_ok=True)
        raise HTTPException(500, f"DOCX导出失败: {e}")
