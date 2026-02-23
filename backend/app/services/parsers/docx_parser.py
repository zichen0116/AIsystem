"""
Word (DOCX) 解析器
提取文本和图片，支持调用视觉模型理解图片内容
"""
from pathlib import Path
from typing import Any, Optional
import uuid
import logging
import os
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import io

from app.services.parsers.base import BaseParser, ParsedChunk, ParseResult
from app.services.ai.vision_service import VisionService

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    """
    Word 文档解析器

    核心功能：
    1. 提取段落文本
    2. 提取表格文本
    3. 提取图片并调用视觉模型生成描述
    """

    def __init__(self, output_dir: str = "media/extracted", api_key: Optional[str] = None):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.api_key = api_key or os.getenv("DASHSCOPE_API_KEY")

        # 初始化视觉服务
        if self.api_key:
            self.vision_service = VisionService(api_key=self.api_key)
        else:
            self.vision_service = None
            logger.warning("未配置 DASHSCOPE_API_KEY，图片将使用基础描述")

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    async def parse(self, file_path: Path) -> ParseResult:
        """
        解析 Word 文档

        Args:
            file_path: DOCX 文件路径

        Returns:
            ParseResult: 包含文本块列表和图片路径列表
        """
        chunks: list[ParsedChunk] = []
        images: list[str] = []

        doc = Document(str(file_path))
        file_name = file_path.name

        # 收集所有文本内容
        all_text = []
        image_refs = []
        image_descriptions = []  # 存储图片描述

        # 1. 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                all_text.append(text)

                # 检查段落中的图片（inline 和 floating 类型）
                for run in para.runs:
                    # Inline 图片
                    for inline in run._element.xpath(".//w:drawing//wp:inline"):
                        try:
                            img_data = await self._extract_and_describe_image(
                                inline, doc, file_name, len(images)
                            )
                            if img_data:
                                img_path, description = img_data
                                images.append(img_path)
                                img_name = Path(img_path).name
                                image_refs.append(f"\n[IMAGE: {img_name}]\n")
                                image_descriptions.append({
                                    "filename": img_name,
                                    "description": description
                                })
                        except Exception as e:
                            logger.error(f"提取段落内嵌图片失败: {e}")

                    # Floating 图片（浮动图片）
                    for anchor in run._element.xpath(".//w:drawing//wp:anchor"):
                        try:
                            img_data = await self._extract_and_describe_image(
                                anchor, doc, file_name, len(images)
                            )
                            if img_data:
                                img_path, description = img_data
                                images.append(img_path)
                                img_name = Path(img_path).name
                                image_refs.append(f"\n[IMAGE: {img_name}]\n")
                                image_descriptions.append({
                                    "filename": img_name,
                                    "description": description
                                })
                        except Exception as e:
                            logger.error(f"提取浮动图片失败: {e}")

        # 2. 提取表格
        for table in doc.tables:
            table_text = self._extract_table_text(table)
            if table_text:
                all_text.append(f"[表格]\n{table_text}")

        # 3. 提取文档中的所有图片（从 relationships）
        try:
            all_doc_images = await self._extract_all_images_from_rels(doc, file_name)
            for img_path, description in all_doc_images:
                if img_path not in images:  # 避免重复
                    images.append(img_path)
                    img_name = Path(img_path).name
                    image_refs.append(f"\n[IMAGE: {img_name}]\n")
                    image_descriptions.append({
                        "filename": img_name,
                        "description": description
                    })
        except Exception as e:
            logger.error(f"提取文档图片失败: {e}")

        # 4. 构建文本块
        if all_text:
            content = "\n".join(all_text)

            # 追加图片引用
            if image_refs:
                content += "\n" + "".join(image_refs)

            chunk = ParsedChunk(
                content=content,
                metadata={
                    "source": file_name,
                    "page": 1,
                    "type": "text",
                    "has_image": len(images) > 0,
                    "image_count": len(images),
                    "image_descriptions": image_descriptions
                }
            )
            chunks.append(chunk)

        logger.info(f"Word 解析完成: {file_name}, 共 {len(chunks)} 个文本块, {len(images)} 张图片")

        return ParseResult(chunks=chunks, images=images)

    async def _extract_and_describe_image(
        self, inline_element, doc: Document, file_name: str, index: int
    ) -> tuple[str, str] | None:
        """提取图片并生成描述"""
        try:
            # 获取图片 relationship ID
            blip = inline_element.xpath(".//a:blip/@r:embed")
            if not blip:
                return None

            rId = blip[0]

            # 通过 relationship 获取图片数据
            image_part = doc.part.related_parts.get(rId)
            if not image_part:
                return None

            # 获取图片尺寸
            extent = inline_element.xpath(".//wp:extent")
            if extent:
                cx = int(extent[0].get("cx", 0)) / 914400  # EMU to inches
                cy = int(extent[0].get("cy", 0)) / 914400
                width, height = int(cx * 96), int(cy * 96)  # inches to pixels
            else:
                width, height = 0, 0

            # 检查是否应该提取
            if not self._should_extract_image(width, height):
                return None

            # 保存图片
            img_data = image_part.blob
            img_ext = image_part.content_type.split('/')[-1] if image_part.content_type else 'png'
            if img_ext == 'jpeg':
                img_ext = 'jpg'

            img_name = f"{file_name}_img{index}_{uuid.uuid4().hex[:8]}.{img_ext}"
            img_path = self.output_dir / img_name

            with open(img_path, 'wb') as f:
                f.write(img_data)

            # 生成图片描述（调用视觉模型）
            description = await self._generate_image_description(str(img_path), width, height)
            return str(img_path), description

        except Exception as e:
            logger.error(f"提取内嵌图片失败: {e}")
            return None

    async def _extract_all_images_from_rels(
        self, doc: Document, file_name: str
    ) -> list[tuple[str, str]]:
        """从所有 relationships 中提取图片"""
        images = []

        try:
            image_index = 0
            for rel_id, rel in doc.part.rels.items():
                # 检查是否是图片类型
                if "image" in rel.target_ref.lower() or "media" in rel.target_ref.lower():
                    try:
                        image_part = rel.target_part
                        img_data = image_part.blob

                        if not img_data:
                            continue

                        # 确定文件扩展名
                        content_type = getattr(image_part, 'content_type', '')
                        if 'png' in content_type:
                            ext = 'png'
                        elif 'jpeg' in content_type:
                            ext = 'jpg'
                        elif 'gif' in content_type:
                            ext = 'gif'
                        elif 'bmp' in content_type:
                            ext = 'bmp'
                        else:
                            ext = rel.target_ref.split('.')[-1] if '.' in rel.target_ref else 'png'
                            if ext not in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
                                ext = 'png'

                        img_name = f"{file_name}_img{image_index}_{uuid.uuid4().hex[:8]}.{ext}"
                        img_path = self.output_dir / img_name

                        with open(img_path, 'wb') as f:
                            f.write(img_data)

                        # 尝试获取图片尺寸
                        try:
                            from PIL import Image
                            with Image.open(io.BytesIO(img_data)) as img:
                                width, height = img.size
                        except:
                            width, height = 0, 0

                        # 生成图片描述
                        if self._should_extract_image(width, height):
                            description = await self._generate_image_description(
                                str(img_path), width, height
                            )
                            images.append((str(img_path), description))
                            image_index += 1
                            logger.info(f"提取图片: {img_path}")

                    except Exception as e:
                        logger.error(f"处理图片 {rel_id} 失败: {e}")
                        continue

        except Exception as e:
            logger.error(f"提取文档图片失败: {e}")

        return images

    async def _generate_image_description(
        self, image_path: str, width: int, height: int
    ) -> str:
        """
        生成图片描述（调用视觉模型或使用基础描述）

        Args:
            image_path: 图片路径
            width: 图片宽度
            height: 图片高度

        Returns:
            str: 图片描述
        """
        # 如果配置了视觉服务，调用视觉模型
        if self.vision_service:
            try:
                prompt = (
                    "请详细分析这张教学图片。提取其中的文字、公式、图表内容，"
                    "并总结画面所传达的知识点。\n\n"
                    "【重要】如果检测到表格，请务必以标准的 Markdown 表格格式提取内容，"
                    "例如：\n"
                    "| 列1 | 列2 | 列3 |\n"
                    "| --- | --- | --- |\n"
                    "| 内容 | 内容 | 内容 |\n\n"
                    "保持数据的结构性，便于后续分析和检索。"
                )
                description = await self.vision_service.describe_image(
                    image_path=image_path,
                    prompt=prompt
                )

                # 检查是否成功
                if description and not description.startswith("错误:") and not description.startswith("API "):
                    return description
                else:
                    logger.warning(f"视觉理解失败，使用基础描述: {description}")
            except Exception as e:
                logger.warning(f"调用视觉服务失败: {e}")

        # 回退到基础描述
        return self._generate_basic_description(width, height)

    def _generate_basic_description(self, width: int, height: int) -> str:
        """生成基础图片描述"""
        aspect_ratio = width / height if height > 0 else 1

        if aspect_ratio > 1.5:
            desc = f"A wide image ({width}x{height}), possibly a chart or diagram"
        elif aspect_ratio < 0.7:
            desc = f"A tall image ({width}x{height})"
        else:
            desc = f"An image ({width}x{height})"

        return desc

    def _extract_table_text(self, table: Table) -> str:
        """提取表格文本"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if cells:
                rows.append(" | ".join(cells))
        return "\n".join(rows)

    def _should_extract_image(self, width: int, height: int, min_size: int = 100) -> bool:
        """判断图片是否应该提取"""
        return width >= min_size and height >= min_size
